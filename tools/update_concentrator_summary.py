from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph

DOCX = Path('/Users/naremovsisyan/Documents/C208-Design/c208design/output/CTE_CompetitiveGameDesign_Concentrator_Outline_26-27.docx')

doc = Document(DOCX)

def replace_text(paragraph, text):
    rpr = deepcopy(paragraph.runs[0]._element.rPr) if paragraph.runs and paragraph.runs[0]._element.rPr is not None else None
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)
    run = paragraph.add_run(text)
    if rpr is not None:
        run._element.insert(0, rpr)

def paragraph_by_text(text):
    return next(p for p in doc.paragraphs if p.text == text)

old_outcome_1 = 'Students will design, program, test, and publish increasingly complex 2D and 3D games in Unity using C#, reusable systems, professional documentation, and iterative production practices.'
old_outcome_2 = 'Students will apply URF Academy design frameworks to multiplayer mechanics, player experience, opposition, interaction, and balance while collaborating in studio teams and building a portfolio aligned to Unity certification and postsecondary pathways.'
outcomes = [
    'Analyze competitive and cooperative game systems using industry vocabulary and URF Academy frameworks for goals, fun, opposition, interaction, and balance.',
    'Apply intermediate C# and Unity skills to create reusable gameplay systems, user interfaces, physics, scene flow, and event-driven interactions.',
    'Create, integrate, and optimize 2D/3D assets, animation, audio, lighting, materials, and visual effects for a cohesive player experience.',
    'Develop a production-ready Game Design Document (GDD), technical plan, task schedule, testing plan, and supporting design documentation.',
    'Prototype, playtest, debug, balance, and revise multiplayer mechanics using qualitative feedback and quantitative evidence.',
    'Collaborate in studio teams to produce, publish, present, and document a polished playable game and individual portfolio artifacts.',
    'Evaluate accessibility, intellectual property, monetization, data privacy, representation, and other ethical responsibilities in game development.',
]

p1 = paragraph_by_text(old_outcome_1)
p2 = paragraph_by_text(old_outcome_2)
replace_text(p1, outcomes[0])
replace_text(p2, outcomes[1])
anchor = p2._element
for outcome in outcomes[2:]:
    new_p = deepcopy(p2._element)
    wrapper = Paragraph(new_p, p2._parent)
    replace_text(wrapper, outcome)
    anchor.addnext(new_p)
    anchor = new_p

old_description = 'This second-year concentrator course advances students from introductory game-design concepts into sustained production. Students use Unity Pathways and C# to build, debug, and publish 2D and 3D games; apply URF Academy workshops to goals, rules, interaction, opposition, game feel, types of fun, and balance; and document decisions through game design documents, technical plans, playtest data, and portfolio reflections. Studio teams rotate through design, programming, art, audio, quality-assurance, and production roles while addressing accessibility, intellectual property, ethics, audience, and industry workflow.'
description_1 = 'This second-year concentrator course advances students from introductory game design into sustained production. Students use Unity Pathways and intermediate C# to design, program, debug, and publish increasingly complex 2D and 3D games. URF Academy design workshops strengthen their understanding of player goals, types of fun, opposition, rules, interaction, game feel, and competitive balance.'
description_2 = 'Through studio-team projects, iterative playtesting, code reviews, industry feedback, and public presentation, students develop professional production habits and explore careers in the games and simulation industry. Students create game design and technical documents, integrate art and audio, address accessibility and ethical responsibilities, prepare portfolio and Unity certification evidence, earn 5 credits per semester, and fulfill the UC a-g “G” elective requirement.'

desc = paragraph_by_text(old_description)
replace_text(desc, description_1)
new_desc = deepcopy(desc._element)
new_desc_wrapper = Paragraph(new_desc, desc._parent)
replace_text(new_desc_wrapper, description_2)
desc._element.addnext(new_desc)

doc.save(DOCX)
print(DOCX)
