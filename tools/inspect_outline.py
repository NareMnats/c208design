from docx import Document
from docx.oxml.ns import qn
import sys

doc = Document(sys.argv[1])
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"P{i:03d} [{p.style.name}] {p.text}")
for ti, table in enumerate(doc.tables):
    print(f"\nTABLE {ti} rows={len(table.rows)} cols={len(table.columns)} style={table.style.name if table.style else None}")
    for ri, row in enumerate(table.rows):
        cells = []
        for ci, cell in enumerate(row.cells):
            text = " | ".join(p.text.replace("\n", " ") for p in cell.paragraphs)
            cells.append(f"C{ci}:{text}")
        print(f"R{ri}: " + " || ".join(cells))
