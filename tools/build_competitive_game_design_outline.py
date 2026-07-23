from copy import deepcopy
from pathlib import Path
from docx import Document

SRC = Path('/Users/naremovsisyan/Downloads/CTE_CompetitiveGameDesign_Movsisyan.N_Outline_26-27.docx')
OUT = Path('/Users/naremovsisyan/Documents/C208-Design/c208design/output/CTE_CompetitiveGameDesign_Concentrator_Outline_26-27.docx')

doc = Document(SRC)

def replace_paragraph(p, text):
    props = deepcopy(p.runs[0]._element.rPr) if p.runs and p.runs[0]._element.rPr is not None else None
    for run in list(p.runs):
        p._element.remove(run._element)
    run = p.add_run(text)
    if props is not None:
        run._element.insert(0, props)

paragraph_updates = {
    19: 'Students will design, program, test, and publish increasingly complex 2D and 3D games in Unity using C#, reusable systems, professional documentation, and iterative production practices.',
    20: 'Students will apply URF Academy design frameworks to multiplayer mechanics, player experience, opposition, interaction, and balance while collaborating in studio teams and building a portfolio aligned to Unity certification and postsecondary pathways.',
    23: 'This second-year concentrator course advances students from introductory game-design concepts into sustained production. Students use Unity Pathways and C# to build, debug, and publish 2D and 3D games; apply URF Academy workshops to goals, rules, interaction, opposition, game feel, types of fun, and balance; and document decisions through game design documents, technical plans, playtest data, and portfolio reflections. Studio teams rotate through design, programming, art, audio, quality-assurance, and production roles while addressing accessibility, intellectual property, ethics, audience, and industry workflow.',
    28: 'Apply intermediate C# programming, event-driven logic, data structures, debugging, and reusable component design in Unity.',
    29: 'Prototype and balance competitive or cooperative multiplayer mechanics using URF Academy design frameworks and evidence from playtests.',
    30: 'Create production-ready game design documents, technical specifications, task boards, source-control routines, and testing plans.',
    31: 'Integrate 2D/3D assets, animation, audio, lighting, user interfaces, physics, and accessibility features into cohesive player experiences.',
    32: 'Collaborate in studio teams to scope, build, test, present, publish, and reflect on a polished game while preparing a portfolio and Unity certification evidence.',
    35: 'Studio-Based Project Learning: Production sprints culminating in playable builds and portfolio evidence.',
    36: 'Unity Pathways: Guided missions in programming, 2D/3D production, UI, audio, animation, lighting, materials, VFX, and publishing.',
    37: 'URF Academy Workshops: Paper prototypes and design labs focused on goals, fun, opposition, rules, interaction, game feel, and balance.',
    38: 'Code Reviews and Technical Demonstrations: Live modeling, pair programming, debugging conferences, and documentation checks.',
    39: 'Iterative Playtesting: Structured observation, data collection, peer critique, accessibility review, and revision.',
    43: 'Industry mentor and guest-speaker critiques of game pitches, portfolios, and production milestones.',
    44: 'Client-style design briefs, game jams, SkillsUSA participation, studio visits, and public exhibition of student work when available.',
    48: 'Software and curriculum: Unity Hub and current supported Unity LTS Editor, Visual Studio Code or Visual Studio, Git-based source control, image/audio creation tools, Unity Learn Pathways, and Riot Games URF Academy curriculum.',
    49: 'Hardware: Development workstation, monitor, keyboard, mouse, headphones/microphone, reliable network access, and optional game controllers or mobile test devices.',
    51: 'Studio Participation: Students are expected to attend consistently, arrive prepared, follow assigned production roles, and contribute evidence of progress during each sprint.',
    52: 'Classwork and Production:',
    53: 'Most learning is completed through demonstrations, Unity Pathway missions, URF Academy design labs, programming exercises, and multiweek studio projects.',
    54: 'Students maintain organized project files, versioned backups, task boards, design documentation, and attribution records.',
    55: 'Playable builds, code, documentation, critique participation, and portfolio artifacts provide the primary evidence of learning.',
    56: 'Teams use checkpoints and peer feedback, but each student is assessed on identifiable individual contributions and technical understanding.',
    57: 'Late or revised work follows classroom procedures and must preserve functional project versions.',
    58: 'Outside-of-class Work:',
    59: 'Projects are planned for completion during class; students may use approved cloud-based learning resources outside class when access permits. Major production deadlines are announced in advance.',
    60: 'Assessments: Skill checks, design critiques, code reviews, playtest reports, Unity Pathway evidence, portfolio reviews, presentations, and a final playable build assess both process and product.',
}

for idx, text in paragraph_updates.items():
    replace_paragraph(doc.paragraphs[idx], text)

units = [
    ('Unit 1', 'Competitive Design Foundations and Studio Workflow', '12', '8',
     'Students revisit core design through URF Academy workshops, analyze competitive and cooperative games, define player goals and types of fun, and establish studio workflow, documentation, scope, file management, and ethical asset use.',
     'D1.1, D1.4, D1.7, D2.1, D2.6', '2.0 (2.5, 2.7), 7.0 (7.4, 7.7), 8.0 (8.3, 8.6)', '1, 2, 8, 10, 11'),
    ('Unit 2', 'C# Systems, Debugging, and Reusable Gameplay', '15', '15',
     'Students build reusable C# components using collections, methods, state, events, input, collisions, scoring, and timers, then document, test, and debug the code.',
     'D7.2, D7.3, D7.4, D7.5', '5.0 (5.2, 5.5, 5.7, 5.8, 5.9), 10.0 (10.1, 10.3, 10.10)', '1, 4, 5, 11'),
    ('Unit 3', 'Multiplayer Mechanics, Opposition, and Balance', '12', '13',
     'Using URF Academy design labs, students create paper and digital multiplayer prototypes; tune rules, resources, feedback loops, risk and reward, fairness, difficulty, and strategic choice using playtest evidence.',
     'D2.6, D2.7, D4.1, D4.3, D4.4, D4.7', '2.0 (2.5), 5.0 (5.3, 5.4, 5.7), 9.0 (9.2, 9.7)', '2, 5, 9, 10, 11'),
    ('Unit 4', 'World Building, Physics, UI, and Accessibility', '12', '13',
     'Students build a cohesive 2D or 3D environment in Unity, implement camera and physics behavior, create responsive user interfaces, and apply HCI, accessibility, performance, and hardware-constraint considerations.',
     'D2.9, D2.10, D5.3, D5.4, D6.3, D6.4', '4.0 (4.1, 4.3), 5.0 (5.3, 5.5), 10.0 (10.3, 10.6)', '1, 4, 5, 10, 12'),
    ('Unit 5', 'Game Feel: Art, Animation, Audio, Lighting, and VFX', '10', '15',
     'Through Unity Pathway missions, students integrate and optimize original or properly licensed visual and audio assets, materials, animation, lighting, particles, feedback, and game-feel techniques for a consistent art direction.',
     'D2.8, D5.1, D5.2, D5.3, D5.4, D7.6', '2.0 (2.6), 8.0 (8.3, 8.6), 10.0 (10.3, 10.6, 10.9)', '4, 8, 10, 12'),
    ('Unit 6', 'AI, NPC Behavior, and Dynamic Challenge', '12', '13',
     'Students design and implement NPC behaviors using finite-state machines, navigation, perception, decision logic, and difficulty tuning, then evaluate whether the behavior supports the intended player experience.',
     'D7.2, D7.3, D8.1, D8.2, D8.3, D8.4', '5.0 (5.3, 5.8, 5.9, 5.10, 5.12), 10.0 (10.1, 10.3)', '1, 4, 5, 10'),
    ('Unit 7', 'Preproduction, Industry Pitch, and Production Planning', '12', '8',
     'Teams respond to a design brief by researching audience and market, defining scope, producing a game design document and technical plan, estimating work, assigning roles, identifying legal and ethical risks, and pitching a production-ready concept.',
     'D1.5, D1.6, D1.7, D1.8, D2.2, D2.3, D3.1, D3.2', '2.0 (2.4, 2.5, 2.7), 3.0 (3.4, 3.5, 3.9), 7.0 (7.2, 7.4, 7.5), 8.0 (8.6)', '2, 3, 7, 8, 9, 11, 12'),
    ('Unit 8', 'Studio Production, Publishing, and Portfolio', '10', '20',
     'Teams complete iterative production sprints, maintain playable builds, conduct formal QA and accessibility testing, resolve defects, publish a final build, present to an authentic audience, and curate individual portfolio and Unity certification evidence.',
     'D3.3, D3.4, D7.2, D7.3, D7.4, D7.5, D7.6', '7.0 (7.3, 7.4, 7.5, 7.7), 9.0 (9.1, 9.2, 9.3, 9.7), 11.0 (11.1, 11.2, 11.5)', '1, 2, 4, 5, 8, 9, 10, 11, 12'),
]

def set_cell(cell, text):
    replace_paragraph(cell.paragraphs[0], text)
    for p in list(cell.paragraphs[1:]):
        cell._element.remove(p._element)

for table, data in zip(doc.tables[7:15], units):
    num, title, ch, lh, desc, pathway, anchor, crp = data
    set_cell(table.cell(0, 0), num)
    set_cell(table.cell(0, 1), title)
    set_cell(table.cell(0, 3), ch)
    set_cell(table.cell(0, 5), lh)
    set_cell(table.cell(1, 0), 'Description: ' + desc)
    set_cell(table.cell(3, 0), pathway)
    set_cell(table.cell(3, 2), anchor)
    set_cell(table.cell(3, 4), crp)

set_cell(doc.tables[15].cell(0, 1), '95')
set_cell(doc.tables[15].cell(1, 1), '105')

competencies = [
    ('Analyze competitive game systems and use URF Academy vocabulary to justify design choices.', 'Establish professional studio workflow, scope, documentation, attribution, and team norms.'),
    ('Write, test, debug, and document reusable C# gameplay components in Unity.', 'Use state, events, collections, input, physics, scoring, and timing to create maintainable systems.'),
    ('Prototype multiplayer goals, rules, opposition, interaction, and feedback loops.', 'Use qualitative and quantitative playtest evidence to improve fairness, depth, and balance.'),
    ('Construct functional 2D/3D environments with physics, cameras, UI, and accessible interaction.', 'Evaluate usability, hardware constraints, and performance when revising the player experience.'),
    ('Integrate art, animation, materials, lighting, VFX, and audio into a coherent game-feel system.', 'Apply copyright, licensing, optimization, and asset-management practices.'),
    ('Implement finite-state-machine or comparable NPC behavior and dynamic challenge.', 'Test whether AI behavior is readable, engaging, and aligned with design goals.'),
    ('Produce and pitch a feasible GDD, technical plan, schedule, risk review, and audience rationale.', 'Perform an assigned studio role while communicating progress and adapting to production constraints.'),
    ('Publish a stable playable build supported by QA evidence and a professional presentation.', 'Curate individual code, design, reflection, certification, and production artifacts in a portfolio.'),
]
for table, pair in zip(doc.tables[16:24], competencies):
    set_cell(table.cell(0, 1), pair[0])
    set_cell(table.cell(1, 1), pair[1])

matrix = [
    ('Unit 1: Competitive Design Foundations and Studio Workflow',
     'Game analysis; player goals; types of fun; ethical asset use; studio roles; project organization.',
     'ELA RSIT 11-12.1 and 11-12.7 - evidence-based analysis.\nCTE D1.4, D1.7, D2.1, D2.6 - player impact, team challenges, design vocabulary, evaluation.',
     'R - Academic Content\nT - Career Tech\nM - Other'),
    ('Unit 2: C# Systems, Debugging, and Reusable Gameplay',
     'C# logic; methods; collections; events; state; reusable components; debugging; code documentation.',
     'Math A-CED.1 and A-CED.2 - model variables and relationships.\nCTE D7.2-D7.5; Anchor 5.5, 5.8, 5.9 - plan, code, test, document, and decompose problems.',
     'R - Academic Content\nT - Career Tech'),
    ('Unit 3: Multiplayer Mechanics, Opposition, and Balance',
     'Paper/digital prototyping; rules; strategic choice; risk/reward; fairness; playtest data; iteration.',
     'Math S-ID.1 and S-ID.2 - represent and compare playtest data.\nCTE D2.6, D2.7, D4.1, D4.3, D4.4, D4.7 - mechanics, players, strategy, action, and challenge.',
     'T - Academic Content\nT - Career Tech\nR - Other'),
    ('Unit 4: World Building, Physics, UI, and Accessibility',
     '2D/3D environments; cameras; physics; UI; HCI; accessibility; constraints; performance.',
     'Math G-MG.1 and G-MG.3 - geometric modeling and design constraints.\nCTE D2.9, D2.10, D5.3, D5.4, D6.3, D6.4 - interface, physics, environment, level, and HCI design.',
     'R - Academic Content\nT - Career Tech\nT - Other'),
    ('Unit 5: Game Feel: Art, Animation, Audio, Lighting, and VFX',
     'Art direction; materials; animation; lighting; particles; audio; feedback; licensing; optimization.',
     'Visual Arts Prof.VA:Cr2.1.II and Cr3.1.II - select, refine, and evaluate artistic work.\nCTE D2.8, D5.1-D5.4, D7.6 - immersion, media integration, tools, environment, and multimedia.',
     'T - Academic Content\nT - Career Tech'),
    ('Unit 6: AI, NPC Behavior, and Dynamic Challenge',
     'Finite-state machines; navigation; perception; decision logic; testing; difficulty tuning.',
     'Math F-IF.4 and F-IF.5 - interpret functions in context.\nCTE D8.1-D8.4; Anchor 5.3, 5.10, 5.12 - AI, intelligent agents, abstraction, and Boolean logic.',
     'R - Academic Content\nT - Career Tech'),
    ('Unit 7: Preproduction, Industry Pitch, and Production Planning',
     'Audience and market research; GDD; technical plan; scope; schedule; roles; ethics; pitch.',
     'ELA WS 11-12.2, 11-12.4, 11-12.6, 11-12.8; SL 11-12.4 - research, technical writing, collaboration, and presentation.\nCTE D1.5-D1.8, D2.2-D2.3, D3.1-D3.2 - industry, lifecycle, documentation, and specification.',
     'T - Academic Content\nT - Career Tech\nR - Other'),
    ('Unit 8: Studio Production, Publishing, and Portfolio',
     'Agile production; version control; QA; accessibility; optimization; publishing; presentation; portfolio.',
     'ELA WS 11-12.5 and SL 11-12.4 - revision and presentation.\nCTE D3.3-D3.4, D7.2-D7.6; Anchor 9.7, 11.2, 11.5 - create, test, present, collaborate, certify, and document work.',
     'T - Academic Content\nT - Career Tech\nT - Other'),
]
mt = doc.tables[24]
for row, data in zip(mt.rows[1:], matrix):
    for cell, text in zip(row.cells, data):
        set_cell(cell, text)

doc.save(OUT)
print(OUT)
